"""Document loaders — extract plain text from uploaded files.

Each loader implements ``load(path) -> LoadedDocument``. Loaders only know how
to turn a file into normalized text (+ page count where available); chunking,
embedding and persistence happen elsewhere.
"""

import logging
from dataclasses import dataclass
from typing import Protocol

import docx
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class LoadedDocument:
    """Text extracted from a document.

    ``page_count`` is ``None`` for formats without a reliable page concept
    (e.g. DOCX).
    """

    text: str
    page_count: int | None


class DocumentLoader(Protocol):
    """Structural type every loader satisfies."""

    def load(self, path: str) -> LoadedDocument: ...


class PDFLoader:
    """Extract text from a PDF using PyPDF2."""

    def load(self, path: str) -> LoadedDocument:
        reader = PdfReader(path)
        pages = reader.pages
        parts: list[str] = []
        for page in pages:
            extracted = page.extract_text() or ""
            if extracted.strip():
                parts.append(extracted)
        text = "\n\n".join(parts)
        logger.debug("PDF loaded: %d pages, %d chars", len(pages), len(text))
        return LoadedDocument(text=text, page_count=len(pages))


class DocxLoader:
    """Extract text from a DOCX using python-docx.

    python-docx has no reliable page count (pagination is computed by the word
    processor), so ``page_count`` is reported as ``None``.
    """

    def load(self, path: str) -> LoadedDocument:
        document = docx.Document(path)
        parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        text = "\n".join(parts)
        logger.debug("DOCX loaded: %d paragraphs, %d chars", len(parts), len(text))
        return LoadedDocument(text=text, page_count=None)
